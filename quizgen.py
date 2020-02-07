"""CM&A Quiz Generator

This generator is creates quiz "packets" from a spreadsheet of 
questions (i.e. "the question database"), such as those from 
the Central Quizzing Leadership Team.

Typical usage:
    import numpy as np
    import quizgen
    np.random.seed(202002081)

    #
    # specify the quizMakeup, i.e. how much from from which books, 
    # chapters, and verses
    #
    quizMakeup={'current':{'frac':0.5,
                           'content':[('HEB',[12,13]),('1P',[1])]},
                'past':{'frac':0.5,
                        'content':[('HEB',[1,2,3,4,5,6,7,8,9,10,11])]}
                }
    nquiz=6

    # 
    # Instantiate the quiz generator with the database
    #
    fnxls='HEB1P2P_CMA_marked.xls'
    QG=quizgen.QuizGenerator(fndatabase=fnxls)

    # add custom limits for certain question types
    QG.quizDistribution['q']['limit']=(150,300)
    QG.quizDistribution['ft']['limit']=(150,300)
    #QG.verbose=True

    # set the quiz makeup and number of quizzes
    QG.setQuizMakeup(quizMakeup,nquiz=nquiz)
    # pull the questions that will be drawn from for this series 
    # of quizzes
    QG.getContent()
    # generate the tabulated questions, and some extras of each 
    # question type
    QG.generateQuizTables(xtra=10)

Additional options:
    QG.verbose = True           # more descriptive logging
    QG.scramblePeriod = False   # keep period blocks in order
    QG.quizType='custom'        # disables distributions

Custom Quizzes
    Custom quizzes can be created by modifying the QG.quizDistribution 
    property such as:
        # make the quiz type 'custom'
        QG.quizType='custom'
        # reset distribution
        # -- to specialize, make the minimums add up the number of 
        #    questions
        # -- to make this more even, make the minimums close to 1/5 
        #    of the total number of questions.
        # -- to be more representative of the question distributions 
        #    provided by CQLT, make the mins/maxes much higher
        qdist=QG.quizDistribution
        qdist['int']['range']=(10,50)
        qdist['cr']['range']=(10,50)
        qdist['ft']['range']=(10,50)
        qdist['ma']['range']=(10,50)
        qdist['q']['range']=(10,50)


Ted Tower, 2/2020
"""
import os
import pandas as pd
import numpy as np
from IPython.display import display
import xlrd
import re
import pprint

# imports from python-docx to create the Word document
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_LINE_SPACING

class QuizGenerator():
    def __init__(self,fndatabase,quizType='epistle',quizDistribution=None):
        print('QuizGenerator initialized.')
        # load database
        self.loadDatabase(fndatabase)

        self.quizType=quizType
        self.nquiz=1
        #self.quizMakeup={'current':{'frac':0.5,'content':[('HEB',[6,7,8,9,10],[150,300])]},
        #                'past':{'frac':0.5,'content':[('HEB',[1,2,3,4,5],[150,300])]}
        #    }

        if(quizDistribution==None):
            # epistle
            qdist={'int':{'range':(9,16),'types':('int',),'label':'Interrogative'},
                'cr':{'range':(3,7),'types':('cr','cvr','cvrma','crma'),'label':'Chapter Reference'},
                'ft':{'range':(3,4),'types':('ft','f2v','ftv','ftn'),'label':'Finish-The-Verse'},
                'ma':{'range':(1,2),'types':('ma',),'label':'Multiple Answer'},
                'q':{'range':(3,4),'types':('q','q2'),'label':'Quote'}}
            self.quizDistribution=qdist
        else:
            self.quizType='custom'
            self.quizDistribution=quizDistribution

        self.quizMakeup=None
        self.verbose=False

        #
        # optional settings
        #
        self.scramblePeriod=1   # set to zero to have questions by period (as a check)

        # inits
        self.quizzes=None
        self.quizStats=None
        self.extraQuestions=None

    def __repr__(self):
        msg="""QuizGenerator instance

    Quiz Type: {quizType}
    Number of quizzes: {nquiz}
    Verbose: {verbose}
    Scramble period: {scramblePeriod}

    Quiz Distribution:
{qdist}

    Quiz Makeup:
{qmakeup}

    
        """.format(quizType=self.quizType,nquiz=self.nquiz,
                    qdist=pprint.pformat(self.quizDistribution,width=60),
                    qmakeup=pprint.pformat(self.quizMakeup,width=30),
                    verbose=self.verbose,
                    scramblePeriod=self.scramblePeriod)


        return(msg)

    def setQuizMakeup(self,quizMakeup,nquiz=1):
        """Set the quiz generator's quizMakeup and nquiz properties."""
        self.quizMakeup=quizMakeup
        self.nquiz=nquiz

    def loadDatabase(self,fnxls):
        """Load the question database

        Args:
            fnxls (string): filename of Excel file
                The expected format of the Excel file is:
                    Book, Chapter, Verse, Verse2, Question, Answer, Group
                    -- Group is added to the database as a flag for 
                    whether it is part of the 150 or 300 key verses.
                    Other group labels are possible.
                    -- This parses the Excel file looking for bolded key 
                    words.  These are extracted internally.

        This function creates a Pandas DataFrame with the following 
        column headings:
            ['BK','CH','VS','VE','TYPE','QUESTION','ANSWER','GROUP',
              'QKEYWORDS','AKEYWORDS','FLAGS','BCV']
        Most of these are straightforward from the Excel file except the 
        following:
            QKEYWORDS, AKEYWORDS -- keywords in the question or answer 
                    (comma separated)
            FLAGS -- currently, this only supports 'repeat'
            BCV   -- a string like <book>_<chapter>_<verse>
                    (e.g. HEB_1_1) to help with not asking another 
                    question that uses the same verse.
        """
        #https://stackoverflow.com/questions/12371787/how-do-i-find-the-formatting-for-a-subset-of-text-in-an-excel-document-cell?rq=1
        # accessing Column 'C' in this example
        COL_IDX = 5

        book = xlrd.open_workbook(fnxls, formatting_info=True)
        sht = book.sheet_by_index(0)

        hdr=[]
        for ii in range(sht.ncols):
            hdr.append(sht.cell_value(0,ii))
        #regcol=list(set(range(sht.ncols)).difference((COL_IDX,)))

        L=[]
        for row_idx in range(1,sht.nrows):
            #if(row_idx>20): break

            # get non-question fields
            row={}
            for ii in range(sht.ncols):
                txt = sht.cell_value(row_idx, ii)
                if(isinstance(txt,str)):
                    txt=txt.replace(u'\xa0', u' ')
                row[hdr[ii]]=txt

            # read question cell and format list
            for COL_IDX in [5,6]:
                text_cell = sht.cell_value(row_idx, COL_IDX)
                text_cell_xf = book.xf_list[sht.cell_xf_index(row_idx, COL_IDX)]

                # skip rows where cell is empty
                if not text_cell:
                    continue
                #print(text_cell)

                text_cell_runlist = sht.rich_text_runlist_map.get((row_idx, COL_IDX))
                if text_cell_runlist:
                    #print(text_cell)
                    #print('(cell multi style) SEGMENTS:')
                    #print(text_cell_runlist)
                    segments = []
                    for segment_idx in range(len(text_cell_runlist)):
                        start = text_cell_runlist[segment_idx][0]
                        # the last segment starts at given 'start' and ends at the end of the string
                        end = None
                        if segment_idx != len(text_cell_runlist) - 1:
                            end = text_cell_runlist[segment_idx + 1][0]
                        segment_text = text_cell[start:end]
                        segments.append({
                            'text': segment_text,
                            'font': book.font_list[text_cell_runlist[segment_idx][1]]
                        })
                        # segments did not start at beginning, assume cell starts with text styled as the cell
                        if text_cell_runlist[0][0] != 0:
                            segments.insert(0, {
                                'text': text_cell[:text_cell_runlist[0][0]],
                                'font': book.font_list[text_cell_xf.font_index]
                            })

                    boldlist=[]
                    for segment in segments:
                        #print('   "%s"'%segment['text'],'italic:',segment['font'].italic,'bold:', segment['font'].bold)
                        if(segment['font'].bold):
                            boldlist.append(segment['text'])
                    keywords=','.join(boldlist)
                else:
                    #print('(cell single style)',
                    keywords=''

                # add question and answer keywords
                if(COL_IDX==5):
                    row['QKEYWORDS']=keywords
                else:
                    row['AKEYWORDS']=keywords

                # add column for flags
                row['FLAGS']=''

                # column for unique verse identifier
                row['BCV']='%s_%d_%d'%(row['BK'],int(row['CH']),int(row['VS']))

            L.append(row)
        
        # make dataframe
        df=pd.DataFrame(L)
        df=df[['BK','CH','VS','VE','TYPE','QUESTION','ANSWER','GROUP','QKEYWORDS','AKEYWORDS','FLAGS','BCV']]
        df = df.astype({'CH': int, 'VS': int})

        self.database=df

    def getContent(self):
        """Get all the content for the quiz in specified range of 
        books & verses.  If a particular question type (e.g. FT) 
        has a limit assigned (e.g. 150), then the GROUP is used 
        to restrict those questions.
        """
        df=self.database
        Q={}
        for period,v in self.quizMakeup.items():
            frames=[]
            #for bk,ch,grp in v['content']:
            for bk,ch in v['content']:
                #print('%s: book %s, ch %s'%(period,bk,str(ch)))
                df1=df[(df['BK']==bk) & df['CH'].isin(ch)]

                #if(len(grp)):
                #    df1=df1[df1['GROUP'].isin(grp)]
                F=[]
                for k,dv in self.quizDistribution.items():
                    f=df1[df1['TYPE'].str.lower().isin(dv['types'])]
                    nrows1=f.shape[0]
                    if('limit' in dv):
                        f=f[f['GROUP'].isin(dv['limit'])]
                    nrows2=f.shape[0]
                    if(nrows2==0):
                        raise Exception('ack!  %d/%d %s questions in content'%(nrows2,nrows1,k))
                    F.append(f)
                df1=pd.concat(F)
                
                frames.append(df1)
            Q[period]=pd.concat(frames)

        self.quizContent=Q

    def pickQuestion(self,dfquiz,dfremaining,qtype=None,BCV=None,questionCounts=None):
        """Pick a question given the current distribution and remaining 
        questions.

        Args:
            dfquiz (dataframe): the current quiz's dataframe
            dfremaining (dataframe): the remaining unpicked questions
            qtype (string): a specified question type (e.g. for extra 
                questions); default: None
            BCV (list): a list of BCVs to exclude
            questionCounts (dict): a dict of counts of each question type
        """

        # initialize question counts
        if(questionCounts==None):
            tcount=self.countTypes(dfquiz)
        else:
            tcount=questionCounts

        # 
        # determine question type (if not specified)
        #
        if(qtype==None):
            nq=dfquiz.shape[0]

            # assume minimum distribution has been met
            minmet=True
            # get type counts
            tcount=self.countTypes(dfquiz)
            for qt,qdata in self.quizDistribution.items():
                # if questionCounts provided, add these to the current counts
                if(questionCounts!=None):
                    tcount[qt]+=questionCounts[qt]
                # if the count is less than the minimum, then the minmet flag is False
                if(tcount[qt]<qdata['range'][0]): minmet=False

            #if(minmet):
            #    print('minimum is met!')
            #else: print('min is NOT met')
            # calc prob of picking each type
            keys=list(self.quizDistribution.keys())
            n=[]
            for qt in keys:
                qdata=self.quizDistribution[qt]
                if(minmet):
                    # if minimum has been met, then prob proportional to num left before max
                    v=qdata['range'][1]-tcount[qt]
                else:
                    # if minimum is not met, then weights determined from questions remaining
                    # yet to be filled to satisfy the minimum
                    v=qdata['range'][0]-tcount[qt]
                n.append(v)

                #print('%s: questions left: %d'%(qt,v))
            weight=[x/sum(n) for x in n]

            # get the question type
            qtpick=np.random.choice(keys,p=weight)
            #print('picked %s'%qtpick)
        else:
            nq=0
            qtpick=qtype

        #
        # get all the questions of this type
        #
        qdata=self.quizDistribution[qtpick]
        df=dfremaining[(dfremaining['TYPE'].str.lower().isin(qdata['types'])) & (dfremaining['used']==0)]
        #print('df orig rows: %d'%df.shape[0])

        # if there is a question in the quiz, exclude book-chapter-verses that are already in the quiz
        if(nq):
            uv=dfquiz['BCV'].unique().tolist()
            if(BCV!=None):
                uv.extend(BCV)
            #print(sorted(uv))
            #print('unique values: %d'%len(uv))
            df=df[~df['BCV'].isin(uv)]
            #print('df rows, that not same BCV: %d'%df.shape[0])

        repeat=False
        if(df.shape[0]==0):
            # if there are no questions in the dataframe, we've run out!  Time to repeat.
            #print('%s repeat!'%qtpick)
            repeat=True
            # pick among all remaining questions of this type
            df=dfremaining[(dfremaining['TYPE'].str.lower().isin(qdata['types']))]
            #print('df rows, w/repeats: %d'%df.shape[0])
        

        # grab one question
        #print('df rows: %d'%df.shape[0])
        q=df.sample(n=1)
        q['used']=1
        if(repeat): q['FLAGS']+='R'
        #display(q)

        # add to current quiz
        dfquiz=pd.concat([dfquiz,q])

        # set this questions to 'used'
        #dfremaining.drop(q.index,inplace=True)
        dfremaining.at[q.index,'used']=1

        return dfquiz,dfremaining

    def countTypes(self,df):
        """get counts of different question types
        """
        tcount={}
        for qt,qdata in self.quizDistribution.items():
            # dataframe of this type of question
            if(df.shape[0]>0):
                dftype=df[df['TYPE'].str.lower().isin(qdata['types'])]
                nrow=dftype.shape[0]
            else:
                nrow=0
            tcount[qt]=nrow
        return tcount


    def genQuiz(self,C,nquestion=30):
        """generate quiz based on content

        Args:
            C (dict): content dict
            nquestion (int): number of questions to be generated
                (default: 30)
        """
        Q1=pd.DataFrame()
        Q2=pd.DataFrame()
        Q3=pd.DataFrame()
        periodCounts={}

        #
        # pick first 20 questions -- these go in the dateframe, Q1
        #
        #    for example
        #        for an "A" quiz, the 50% of the questions are picked from the "past" period,
        #                         then 50% of the questions from the "current" period
        for period,dfremaining in C.items():
            # pick some questions from this period
            nq=int(20*self.quizMakeup[period]['frac'])
            if(self.verbose): print('picking first %d questions from %s'%(nq,period))
            for qi in range(nq):
                #print('...%d'%qi)
                #print('pick question from %s'%period)
                Q1,dfremaining=self.pickQuestion(Q1,dfremaining)
                #display('count: %d'%dfremaining['used'].value_counts().iloc[1])

            # add the period counts
            periodCounts[period]=[nq]

        if(self.verbose):
            # show the distribution for each question type
            for qt,cnt in self.countTypes(Q1).items():
                rng=self.quizDistribution[qt]['range']
                #print('%s: %d (min: %d, max: %d)'%(qt,cnt,rng[0],rng[1]))
        
                #print('question(%s) %d: %d questions (%d remaining)'%(period,qi,Q1.shape[0],dfremaining.shape[0]))
        
        # scramble first questions
        if(self.scramblePeriod):
            Q1=Q1.sample(frac=1.0)
        # label first questions
        nq1=Q1.shape[0]
        lbl=[str(x+1) for x in range(nq1)]
        Q1['qn']=lbl

        #
        # pick for rest of quiz (16AB, 17AB, 18AB, 19AB, 20AB) -- these go in dataframe, Q2
        #
        if(self.quizType!='custom'):
            nq2=10
        else:
            nq2=nquestion-20
        usedVerses=Q1['BCV'].unique().tolist()
        q1counts=self.countTypes(Q1)
        #pprint.pprint(q1counts)

        for period,dfremaining in C.items():
            nq=int(nq2*self.quizMakeup[period]['frac']+1)
            if(self.verbose): print('picking second %d questions from %s'%(nq,period))
            for qi in range(nq):
                Q2,dfremaining=self.pickQuestion(Q2,dfremaining,BCV=usedVerses,questionCounts=q1counts)
                if(Q2.shape[0]>=nq2): break

            periodCounts[period].append(nq)

        #pprint.pprint(q1counts)

        if(self.verbose):
            # show the distribution for each question type
            for qt,cnt in self.countTypes(Q2).items():
                rng=self.quizDistribution[qt]['range']
                q12cnt=q1counts[qt]+cnt
                #print('%s: %d (Q1: %d, Q2: %d, min: %d, max: %d)'%(qt,q12cnt,q1counts[qt],cnt,rng[0],rng[1]))

        # scramble second part
        if(self.scramblePeriod):
            Q2=Q2.sample(frac=1.0)
        # label second questions
        if(self.quizType!='custom'):
            lbl=['16A','16B','17A','17B','18A','18B','19A','19B','20A','20B']
        else:
            nq2=Q2.shape[0]
            lbl=[str(x+1+nq1) for x in range(nq2)]
        Q2['qn']=lbl


        #display(Q2)

        #
        # pick a few overtime questions
        #
        if(self.quizType!='custom'):
            nq3=3
        else:
            nq3=0
        usedVerses=Q1['BCV'].unique().tolist()
        usedVerses2=Q2['BCV'].unique().tolist()
        usedVerses.extend(usedVerses2)
        #print(usedVerses)
        q12counts=self.countTypes(Q1)
        for qt,cnt in self.countTypes(Q2).items():
            q12counts[qt]+=cnt
        
        #pprint.pprint(q12counts)

        for period,dfremaining in C.items():
            nq=int(nq3*self.quizMakeup[period]['frac']+1)
            if(self.quizType=='custom'): nq=0

            if(self.verbose): print('picking third %d questions from %s'%(nq,period))
            for qi in range(nq):
                #print(list(self.quizDistribution.keys()))
                #
                # for overtime, randomly pick a question type
                #
                qt=np.random.permutation(list(self.quizDistribution.keys()))[0]
                #print(qt)
                #print(qt[0])
                Q3,dfremaining=self.pickQuestion(Q3,dfremaining,BCV=usedVerses,qtype=qt)
                if(Q3.shape[0]>=nq3): break

        # scramble third part
        #Q2.sample(frac=1.0)
        # label overtime questions
        if(self.quizType!='custom'):
            #lbl=['21','21A','21B','22','22A','22B','23','23A','23B']
            lbl=['21','22','23']
        else:
            nq3=Q3.shape[0]
            lbl=[str(x+1+nq1+nq2) for x in range(nq3)]
        Q3['qn']=lbl

        #display(Q3)
        #
        # reorder questions
        #
        rng=['1','2','3','4','5','6','7','8','9','10','11','12','13','14','15','16',
                '16A','16B','17','17A','17B','18','18A','18B','19','19A','19B','20','20A','20B',
                '21','22','23']
        frames=[]
        if(self.quizType!='custom'):
            for r in rng:
                df=Q1[Q1['qn']==r]
                frames.append(df)
                df=Q2[Q2['qn']==r]
                frames.append(df)
                df=Q3[Q3['qn']==r]
                frames.append(df)
        else:
            frames.append(Q1)
            frames.append(Q2)
            #frames.append(Q3)
        
        dfq=pd.concat(frames,sort=False)

        #display(dfq)
        #1/0
        #
        # uncomment this to write out sorted by verse
        #
        #dfq=dfq.sort_values(by=['CH','VS'])

        stats={'min':self.countTypes(Q1),
               'max':q12counts,
               'period':periodCounts}

        return dfq,C,stats

    def genExtraQuestions(self,C,qtype,xtra):
        """pick extra questions
        """
        Q1=pd.DataFrame()
        for period,dfremaining in C.items():
            nq=int(xtra*self.quizMakeup[period]['frac'])
            for qi in range(nq):
                Q1,dfremaining=self.pickQuestion(Q1,dfremaining,qtype=qtype)
                if(Q1.shape[0]>=xtra): break
        # scramble first questions
        Q1.sample(frac=1.0)

        nq1=Q1.shape[0]
        lbl=[str(x+1) for x in range(nq1)]
        Q1['qn']=lbl
        return (Q1,C)

    def generateQuizTables(self,xtra=5,nquestion=30):
        """Generate quiz tables

        Args:
            xtra (int): number of extra questions to generated
            nquestion (int): number of questions to generate per quiz
        """

        # get copy of content from each period
        #    typically, C={'past':dataFrame,'current':dataFrame}
        C={}
        for period,df in self.quizContent.items():
            df['used']=0
            C[period]=df.copy()
        
        # loop through requested quizzes
        QQ=[];QQstats=[]
        for qi in range(self.nquiz):
            dfq,C,stats=self.genQuiz(C,nquestion=nquestion)
            QQ.append(dfq)
            QQstats.append(stats)

        qxtra={}
        for qt,qdata in self.quizDistribution.items():
            Q1,C=self.genExtraQuestions(C,qt,xtra)
            qxtra[qt]=Q1

        self.quizzes=QQ
        self.quizStats=QQstats
        self.extraQuestions=qxtra

    def boldText(self, cell, text, keywords):
        """boldText"""
        if(len(keywords)<1 or (len(keywords[0])<1)):
            p=cell.paragraphs[-1]
            p.add_run(text)
        else:
            keywords=list(set(keywords))
            nk=len(keywords)
            # get start/len for keywords
            #print(text)
            IL=[]
            for k in keywords:
                #print('KEYWORD: %s'%k)
                #IL.append((text.index(k),len(k)))
                rei=re.finditer(k,text)
                for m in rei:
                    k=m.group()
                    startidx=m.span()[0]
                    #print('keyword: %s, startidx: %d'%(k,startidx))
                    IL.append((k,startidx))
            
            # sort according to starting position
            IL=sorted(IL,key=lambda x:x[1])
            #print('IL: %s'%str(IL))

            # start
            txt='%s'%text[:IL[0][1]]
            #txt='something'
            p=cell.paragraphs[-1]
            p.add_run(txt)
            
            for ii,il in enumerate(IL):
                kw,startidx=il
                # bold keyword
                p.add_run(kw).bold = True

                # text in between keywords
                #start=IL[ii][0]+IL[ii][1]
                start=startidx+len(kw)
                end=None
                if(ii<(len(IL)-1)):
                    #end=IL[ii+1][0]
                    end=IL[ii+1][1]
                txt=text[start:end]
                p.add_run(txt).bold = False



    def genQuizPacket(self,fn,title='Quizzes',msg=None):
        """This method creates the quiz packet Word document.
        
         Args:
           fn (string): output filename of quiz
           title (string): title in the document
           msg (string): optional message to write
        """
        

        #
        # document, paragraph, section, font settings
        #
        # -- width for columns: question number, type, Q/A, reference
        width=[Inches(0.375),Inches(0.375),Inches(5.25),Inches(1.)]

        document = Document()
        sections = document.sections
        section = sections[0]
        section.left_margin = Inches(0.75)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9)

        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_after = Pt(3)

        #
        # add the title
        #
        document.add_heading(title, 0)

        #
        # add the message
        #
        if(msg==None):
            if(self.quizType!='custom'):
                msg={'intro':'This is a quiz packet for WGLD.  The quiz packet should have these characteristics:',
                    'list':['Unique questions for each quiz (if possible) in the packet',
                            'Satisfaction of question minimums and maximums for each type',
                            '"A" division quizzes have 50% current and 50% past periods.',
                            '"B" division quizzes are only current content, and will therefore have repeats.'+\
                            '  We have tried to keep these in the alternative questions 16A, 16B, etc.  Replace as necessary.']}
            else:
                msg={'intro':'This is a custom quiz.'}
        
        p = document.add_paragraph(msg['intro'])

        #document.add_paragraph('Unique questions for each quiz (if possible) in the packet', style='List Bullet')
        #document.add_paragraph(
        #    'Satisfaction of question minimums and maximums for each type', style='List Bullet'
        #)
        if('list' in msg):
            for m in msg['list']:
                document.add_paragraph(m, style='List Bullet')

        #
        # loop through quizzes
        #
        for qi,QZ in enumerate(self.quizzes):
            if(qi>0):
                document.add_page_break()
            document.add_heading('Quiz %d'%(qi+1), 1)

            table = document.add_table(rows=1, cols=4)
            #table.style = 'LightShading-Accent1'
            table.style = 'LightGrid-Accent1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Type'
            hdr_cells[2].text = 'Question'
            hdr_cells[3].text = 'Verse'
            for k,cell in enumerate(hdr_cells):
                cell.width=width[k]

            #
            # loop through questions
            #
            ii=0
            for idx,row in QZ.iterrows():
                ii+=1
                row_cells = table.add_row().cells

                # Question Number
                row_cells[0].text = row.qn
                # Question Type
                row_cells[1].text = row.TYPE

                # https://stackoverflow.com/questions/36894424/creating-a-table-in-python-docx-and-bolding-text#36897305

                #
                # Question/Answer cell
                #
                c=row_cells[2]
                q='Q: %s'%row.QUESTION
                keywords=row.QKEYWORDS.split(',')

                self.boldText(cell=c, text=q, keywords=keywords)
                c.add_paragraph()

                #
                # ANSWER
                #
                a='A: %s'%row.ANSWER
                keywords=row.AKEYWORDS.split(',')
                self.boldText(cell=c, text=a, keywords=keywords)
                # book, chapter, verse, and group (e.g. 150,300)
                txt='%s %s:%s'%(row.BK,row.CH,row.VS)
                if(isinstance(row.VE,float)):
                    txt+='-%s'%str(int(row.VE))
                if(isinstance(row.GROUP,float)):
                    txt+='\n(%d)'%row.GROUP
                
                # additional flags (repeats)
                c=row_cells[3]
                if('R' in row['FLAGS']):
                    txt+='\nrepeat'
                    c._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w'))))
                c.text=txt

                # adjust width
                for k,cell in enumerate(row_cells):
                    cell.width=width[k]

            #
            # quiz stats
            #
            stats=self.quizStats[qi]
            qdist=self.quizDistribution
            if(self.quizType!='custom'):
                #
                # normal quiz
                #
                # -- min distribution
                msg='Regular quiz distribution (does not include overtime); ';first=1
                # loop through all the types to show minimums
                for qt,cnt in stats['min'].items():
                    msg+='%s:%d-%d ('%(qt.upper(),cnt,stats['max'][qt])
                    if(first):
                        first=0;
                        msg+='req: '
                    msg+='%d-%d), '%(qdist[qt]['range'][0],qdist[qt]['range'][1])
                msg=msg[:-2]   # get rid of trailing space and comma at end
                #
                # -- period stats
                #
                msg+='; Question counts by period (numbered): '
                for period,cnts in stats['period'].items():
                    msg+='%s=%d; '%(period,cnts[0])
                msg=msg[:-2]
            else:
                # custom quiz
                msg='Custom quiz distribution; '
                for qt,cnt in self.countTypes(QZ).items():
                    msg+=' %s(%d),'%(qt,cnt)
                msg=msg[:-1]
                print(msg)
            #
            # add stats to document
            #
            document.add_paragraph(msg)
        
        #
        # extra question
        #
        if(self.quizType!='custom'):
            document.add_page_break()
            document.add_heading('Extra Questions', level=1)

            msg="""This section contains extra questions of each type for use during the quiz day.
            Make sure to mark the questions used as you use them.
            """
            p = document.add_paragraph(msg)

            for qt,v in self.extraQuestions.items():
                tlist=', '.join([x.upper() for x in self.quizDistribution[qt]['types']])
                document.add_heading('%s Extra Questions (%s)'%(self.quizDistribution[qt]['label'],tlist), level=2)
                
                table = document.add_table(rows=1, cols=4)
                table.style = 'LightGrid-Accent1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '#'
                hdr_cells[1].text = 'Type'
                hdr_cells[2].text = 'Question'
                hdr_cells[3].text = 'Verse'
                for k,cell in enumerate(hdr_cells):
                    cell.width=width[k]
                    
                ii=0
                for idx,row in self.extraQuestions[qt].iterrows():
                    ii+=1
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(ii)
                    row_cells[0].width=width[0]
                    row_cells[1].text = row.TYPE
                    #row_cells[2].text = 'Q: %s\n\nA: %s'%(row.QUESTION,row.ANSWER)

                    #
                    # QUESTION
                    #
                    c=row_cells[2]
                    q='Q: %s'%row.QUESTION
                    keywords=row.QKEYWORDS.split(',')
                    self.boldText(cell=c, text=q, keywords=keywords)

                    c.add_paragraph()
                    #c.add_paragraph()

                    #
                    # ANSWER
                    #
                    a='A: %s'%row.ANSWER
                    keywords=row.AKEYWORDS.split(',')
                    self.boldText(cell=c, text=a, keywords=keywords)

                    #
                    # VERSES
                    #
                    txt='%s %s:%s'%(row.BK,row.CH,row.VS)
                    if(isinstance(row.VE,float)):
                        txt+='-%s'%str(int(row.VE))
                    if(isinstance(row.GROUP,float)):
                        txt+='\n(%d)'%row.GROUP

                    #if(not np.isnan(row.VE)):
                    #    txt+='-%s'%row.VE
                    #row_cells[3].text = txt
                    c=row_cells[3]
                    if('R' in row['FLAGS']):
                        txt+='\nrepeat'
                        c._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w'))))
                    c.text = txt

                    for k,cell in enumerate(row_cells):
                        cell.width=width[k]

        document.save(fn)
        print('Done writing quiz packet (%s)'%fn)

if(__name__=='__main__'):
    print('quiz gen is not meant to be called directly')