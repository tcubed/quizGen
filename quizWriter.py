
import re

# imports from python-docx to create the Word document
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_LINE_SPACING

import logging
#logging.basicConfig(filename='quiz_generator.log',level=logging.DEBUG)
# create logger
logger = logging.getLogger('quiz_writer')
logger.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s|%(name)s|%(levelname)s|%(funcName)s|%(message)s')
logger.handlers=[]
# create console handler and set level to debug
ch = logging.StreamHandler()
ch.setLevel(logging.INFO)
ch.setFormatter(formatter)
logger.addHandler(ch)
# file handler
fh = logging.FileHandler('quiz_writer.log')
fh.setFormatter(formatter)
logger.addHandler(fh)

def countTypes(df,quizDistribution):
    """get counts of different question types
    """
    tcount={}
    for qt,qdata in quizDistribution.items():
        # dataframe of this type of question
        if(df.shape[0]>0):
            if(qt=='sit'):
                dftype=df[df['TYPE'].str.lower().str.startswith(qt)]
            else:
                dftype=df[df['TYPE'].str.lower().isin(qdata['types'])]
            nrow=dftype.shape[0]
        else:
            nrow=0
        tcount[qt]=nrow
    return tcount

class QuizWriter():
    def __init__(self):
        self.loose=False
        pass
        
    def boldText(self, cell, text, keywords):
        """boldText"""
        if(len(keywords)<1 or (len(keywords[0])<1)):
            p=cell.paragraphs[-1]
            p.add_run(text)
        else:
            keywords=list(set(keywords))
            nk=len(keywords)
            # get start/len for keywords
            #print(text)
            IL=[]
            for k in keywords:
                #print('KEYWORD: %s'%k)
                #IL.append((text.index(k),len(k)))
                try:
                    rei=re.finditer(k,text)
                except:
                    print('what?')
                for m in rei:
                    k=m.group()
                    startidx=m.span()[0]
                    #print('keyword: %s, startidx: %d'%(k,startidx))
                    IL.append((k,startidx))
            
            # sort according to starting position
            IL=sorted(IL,key=lambda x:x[1])
            #print(keywords)
            #print('IL: %s'%str(IL))
            #if('paths' in keywords):
            #    print('IL: %s'%str(IL))

            # start
            txt='%s'%text[:IL[0][1]]
            #txt='something'
            p=cell.paragraphs[-1]
            p.add_run(txt)
            
            for ii,il in enumerate(IL):
                kw,startidx=il
                # bold keyword
                p.add_run(kw).bold = True

                # text in between keywords
                #start=IL[ii][0]+IL[ii][1]
                start=startidx+len(kw)
                end=None
                if(ii<(len(IL)-1)):
                    #end=IL[ii+1][0]
                    end=IL[ii+1][1]
                txt=text[start:end]
                p.add_run(txt).bold = False



    def save(self,fn,quizData,title='CMA Bible Quizzes',msg=None):
        """This method creates the quiz packet Word document.
        
         Args:
           fn (string): output filename of quiz
           quizData (dict): quizData object
           title (string): title in the document
           msg (string): optional message to write
        """
        

        #
        # document, paragraph, section, font settings
        #
        # -- width for columns: question number, type, Q/A, reference
        if(quizData['type']=='epistle'):
            width=[Inches(0.375),Inches(0.375),Inches(5.25),Inches(1.)]
        else:
            width=[Inches(0.375),Inches(1),Inches(4.625),Inches(1.)]

        document = Document()
        sections = document.sections
        section = sections[0]
        section.left_margin = Inches(0.75)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9)

        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_after = Pt(3)

        #
        # add the title
        #
        #if(self.loose):
        loose=False
        for qd in quizData['stats']:
            loose=loose or qd['loose']
        if(loose):
            title='%s (loose)'%title
        document.add_heading(title, 0)

        #
        # add the message
        #
        # if(msg==None):
        #     if(self.quizType!='custom'):
        #         msg={'intro':'This is a quiz packet for WGLD.  The quiz packet should have these characteristics:',
        #             'list':['Unique questions for each quiz (if possible) in the packet',
        #                     'Satisfaction of question minimums and maximums for each type',
        #                     '"A" division quizzes have 50% current and 50% past periods.',
        #                     '"B" division quizzes are only current content, and will therefore have repeats.'+\
        #                     '  We have tried to keep these in the alternative questions 16A, 16B, etc.  Replace as necessary.']}
        #     else:
        #         msg={'intro':'This is a custom quiz.'}
        
        # p = document.add_paragraph(msg['intro'])

        # #document.add_paragraph('Unique questions for each quiz (if possible) in the packet', style='List Bullet')
        # #document.add_paragraph(
        # #    'Satisfaction of question minimums and maximums for each type', style='List Bullet'
        # #)
        # if('list' in msg):
        #     for m in msg['list']:
        #         document.add_paragraph(m, style='List Bullet')

        # default message
        if(msg==None):
            msg=[{'type':'p','text':'This is a CM&A quiz packet.  The quiz packet should have these characteristics:'},
                 {'type':'list','text':['Unique questions for each quiz (if possible) in the packet',
                            'Satisfaction of question minimums and maximums for each type (2018 rules)',
                            '"A" division quizzes have 50% current and 50% past periods.',
                            '"B" division quizzes are only current content, and will therefore have repeats.'+\
                            '  We have tried to keep these in the alternative questions 16A, 16B, etc.  Replace as necessary.']}]

        for ii,m in enumerate(msg):
            if(m['type']=='p'):
                # normal paragraph
                p = document.add_paragraph(m['text'])
            elif(m['type']=='list'):
                # bulleted list
                for mitem in m['text']:
                    document.add_paragraph(mitem, style='List Bullet')

        #
        # loop through quizzes
        #
        for qi,QZ in enumerate(quizData['quizzes']):

            chapList=sorted(QZ['CH'].unique())
            logger.debug('chapters: %s'%str(chapList))
            
            stats=quizData['stats'][qi]

            if(qi>0):
                document.add_page_break()
            heading='Quiz %d'%(qi+1)
            if(stats['loose']):
                heading+=' (loose)'
            document.add_heading(heading, 1)

            table = document.add_table(rows=1, cols=4)
            #table.style = 'LightShading-Accent1'
            table.style = 'LightGrid-Accent1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Type'
            hdr_cells[2].text = 'Question'
            hdr_cells[3].text = 'Verse'
            for k,cell in enumerate(hdr_cells):
                cell.width=width[k]

            #
            # loop through questions
            #
            ii=0
            for idx,row in QZ.iterrows():
                ii+=1
                row_cells = table.add_row().cells

                # Question Number
                row_cells[0].text = row.qn
                # Question Type
                row_cells[1].text = row.TYPE

                # https://stackoverflow.com/questions/36894424/creating-a-table-in-python-docx-and-bolding-text#36897305

                #
                # Question/Answer cell
                #
                c=row_cells[2]
                q='Q: %s'%row.QUESTION
                if('QKEYWORDS' in row):
                    keywords=row.QKEYWORDS.split(',')
                else:
                    keywords=[]
                self.boldText(cell=c, text=q, keywords=keywords)
                c.add_paragraph()

                #
                # ANSWER
                #
                a='A: %s'%row.ANSWER
                if('AKEYWORDS' in row):
                    keywords=row.AKEYWORDS.split(',')
                else:
                    keywords=[]
                self.boldText(cell=c, text=a, keywords=keywords)

                # book, chapter, verse, and club (e.g. 150,300)
                txt='%s %s:%s'%(row.BK,row.CH,row.VS)
                # if('VE' in row):
                #     if(isinstance(row.VE,float)):
                #         txt+='-%s'%str(int(row.VE))
                if('2' in row.TYPE):
                    txt+='-%s'%str(int(row.VS)+1)
                
                txt+='\n('
                #if(isinstance(row.CLUB,float)):
                if(len(row.CLUB)):
                    #txt+='\n(%d)'%row.CLUB
                    txt+='%s,'%row.CLUB
                if(row.SET is not None):
                    txt+='%s'%row.SET
                txt+=')'
                
                # additional flags (repeats)
                c=row_cells[3]
                if('R' in row['FLAGS']):
                    txt+='\nrepeat'
                    c._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w'))))
                c.text=txt

                # adjust width
                for k,cell in enumerate(row_cells):
                    cell.width=width[k]

            #
            # quiz stats
            #
            
            qdist=quizData['distribution']
            if(quizData['type']!='custom'):
                #
                # normal quiz
                #
                # -- min distribution
                msg='Quiz distribution (<1-20 only>-<total w/AB>; not including overtime); ';first=1
                # loop through all the types to show minimums
                for qt,cnt in stats['min'].items():
                    msg+='%s:%d-%d ('%(qt.upper(),cnt,stats['max'][qt])
                    if(first):
                        first=0;
                        msg+='required: '
                    msg+='%d-%d), '%(qdist[qt]['range'][0],qdist[qt]['range'][1])
                msg=msg[:-2]   # get rid of trailing space and comma at end
                #
                # -- period stats
                #
                msg+='; Question counts by period (numbered): '
                for period,cnts in stats['period'].items():
                    msg+='%s=%d; '%(period,cnts[0])
                msg=msg[:-2]
            else:
                # custom quiz
                msg='Custom quiz distribution; '
                #for qt,cnt in self._countTypes(QZ).items():
                for qt,cnt in countTypes(QZ,qdist).items():
                    msg+=' %s(%d),'%(qt,cnt)
                msg=msg[:-1]
                print(msg)
            #
            # add stats to document
            #
            document.add_paragraph(msg)
        
        #
        # extra question
        #
        if(quizData['type']!='custom'):
            document.add_page_break()
            document.add_heading('Extra Questions', level=1)

            msg="""This section contains extra questions of each type for use during the quiz day.
            Make sure to mark the questions used as you use them.
            """
            p = document.add_paragraph(msg)

            for qt,v in quizData['extraQuestions'].items():
                tlist=', '.join([x.upper() for x in quizData['distribution'][qt]['types']])
                document.add_heading('%s Extra Questions (%s)'%(quizData['distribution'][qt]['label'],tlist), level=2)
                
                table = document.add_table(rows=1, cols=4)
                table.style = 'LightGrid-Accent1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '#'
                hdr_cells[1].text = 'Type'
                hdr_cells[2].text = 'Question'
                hdr_cells[3].text = 'Verse'
                for k,cell in enumerate(hdr_cells):
                    cell.width=width[k]
                    
                ii=0
                for idx,row in quizData['extraQuestions'][qt].iterrows():
                    ii+=1
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(ii)
                    row_cells[0].width=width[0]
                    row_cells[1].text = row.TYPE
                    #row_cells[2].text = 'Q: %s\n\nA: %s'%(row.QUESTION,row.ANSWER)

                    #
                    # QUESTION
                    #
                    c=row_cells[2]
                    q='Q: %s'%row.QUESTION
                    if('QKEYWORDS' in row):
                        keywords=row.QKEYWORDS.split(',')
                    else:
                        keywords=[]
                    self.boldText(cell=c, text=q, keywords=keywords)

                    c.add_paragraph()
                    #c.add_paragraph()

                    #
                    # ANSWER
                    #
                    a='A: %s'%row.ANSWER
                    if('AKEYWORDS' in row):
                        keywords=row.AKEYWORDS.split(',')
                    else:
                        keywords=[]
                    self.boldText(cell=c, text=a, keywords=keywords)

                    #
                    # VERSES
                    #
                    txt='%s %s:%s'%(row.BK,row.CH,row.VS)
                    #if('VE' in row):
                    #    if(isinstance(row.VE,float)):
                    #        txt+='-%s'%str(int(row.VE))
                    if('2' in row.TYPE):
                        txt+='-%s'%str(int(row.VS)+1)
                    
                    #if(isinstance(row.CLUB,float)):
                    #    txt+='\n(%d)'%row.CLUB
                    if(len(row.CLUB)):
                        txt+='\n(%s)'%row.CLUB
                    
                    #if(not np.isnan(row.VE)):
                    #    txt+='-%s'%row.VE
                    #row_cells[3].text = txt
                    c=row_cells[3]
                    if('R' in row['FLAGS']):
                        txt+='\nrepeat'
                        c._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w'))))
                    c.text = txt

                    for k,cell in enumerate(row_cells):
                        cell.width=width[k]

        document.save(fn)
        print('Done writing quiz packet (%s)'%fn)

        